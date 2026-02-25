"""
Generate 12_entity_summary.docx — Entity Documentation Summary for
Sunshine Palms Holdings LLC in connection with the sale of
2750 Malabar Road SE, Palm Bay, FL 32907.
"""
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from data import *

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
NAVY_RGB = RGBColor(0x1E, 0x3A, 0x5F)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
FONT_NAME = "Calibri"


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_run_font(run, name=FONT_NAME, size=11, bold=False, italic=False,
                 color=None):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure the font name is respected by Word
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text, bold=False, size=11, alignment=None,
                  space_after=6, space_before=0, color=None, italic=False):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_section_heading(doc, number, title):
    """Add a section heading like 'Section 1 - Entity Information'."""
    p = doc.add_paragraph()
    run = p.add_run(f"Section {number} \u2014 {title}")
    set_run_font(run, bold=True, size=12, color=NAVY_RGB)
    pf = p.paragraph_format
    pf.space_before = Pt(16)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    return p


def add_body_text(doc, text, indent=0.0, space_after=4, bold=False,
                  italic=False):
    """Add a regular body text paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold, italic=italic)
    pf = p.paragraph_format
    if indent > 0:
        pf.left_indent = Inches(indent)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(2)
    return p


def add_bullet(doc, text, indent=0.5):
    """Add a bullet-style paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"\u2022  {text}")
    set_run_font(run, size=11)
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.space_after = Pt(2)
    pf.space_before = Pt(1)
    return p


def add_sub_bullet(doc, text, indent=0.75):
    """Add a sub-bullet (dash) paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"\u2013  {text}")
    set_run_font(run, size=11)
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.space_after = Pt(2)
    pf.space_before = Pt(1)
    return p


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, color="000000", size="4"):
    """Set thin borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def format_header_cell(cell, text):
    """Format a table header cell with navy background and white text."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=10, bold=True, color=WHITE_RGB)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_cell_shading(cell, NAVY)
    set_cell_borders(cell, color=NAVY)


def format_body_cell(cell, text, bold=False):
    """Format a table body cell with standard styling."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, size=10, bold=bold)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_cell_borders(cell)


def create_two_col_table(doc, rows_data, col_widths=(2.5, 4.5)):
    """Create a two-column table (Field | Value) with navy headers."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])

    # Header row
    format_header_cell(table.rows[0].cells[0], "Field")
    format_header_cell(table.rows[0].cells[1], "Value")

    # Data rows
    for i, (field, value) in enumerate(rows_data):
        format_body_cell(table.rows[i + 1].cells[0], field, bold=True)
        format_body_cell(table.rows[i + 1].cells[1], value)

    return table


def create_three_col_table(doc, headers, rows_data, col_widths=(2.5, 1.5, 3.0)):
    """Create a three-column table with navy headers."""
    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Set column widths
    for row in table.rows:
        for j, w in enumerate(col_widths):
            row.cells[j].width = Inches(w)

    # Header row
    for j, h in enumerate(headers):
        format_header_cell(table.rows[0].cells[j], h)

    # Data rows
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            format_body_cell(table.rows[i + 1].cells[j], val)

    return table


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_entity_summary():
    doc = Document()

    # -- Default font for the whole document --------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)

    # Set 1" margins all around
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===================================================================
    # TITLE
    # ===================================================================
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("ENTITY DOCUMENTATION SUMMARY")
    set_run_font(title_run, size=14, bold=True, color=NAVY_RGB)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)

    # Subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run(
        f"Prepared for Sale of {PROPERTY['address']}"
    )
    set_run_font(subtitle_run, size=12, bold=False, italic=True)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(6)

    # Horizontal rule (thin line)
    hr_p = doc.add_paragraph()
    hr_p.paragraph_format.space_after = Pt(12)
    pPr = hr_p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{NAVY}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ===================================================================
    # SECTION 1 — Entity Information
    # ===================================================================
    add_section_heading(doc, "1", "Entity Information")

    entity_info = [
        ("Entity Name", PROPERTY["owner_entity"]),
        ("Entity Type", "Limited Liability Company"),
        ("State of Formation", "Florida"),
        ("Document Number", "L18000045678"),
        ("Filing Date", "February 1, 2018"),
        ("Status", "Active"),
        ("Last Annual Report", "Filed through 2026"),
        ("FEI/EIN", "82-XXXXXXX"),
    ]
    create_two_col_table(doc, entity_info)

    # ===================================================================
    # SECTION 2 — Registered Agent
    # ===================================================================
    add_section_heading(doc, "2", "Registered Agent")

    add_bullet(doc, "Name: Mariam Shapira")
    add_bullet(doc, "Address: 123 Ocean Ave, Satellite Beach, FL 32937")
    add_bullet(doc, "Status: Current")

    # ===================================================================
    # SECTION 3 — Members/Managers
    # ===================================================================
    add_section_heading(doc, "3", "Members/Managers")

    create_three_col_table(
        doc,
        headers=["Member", "Ownership", "Role"],
        rows_data=[
            ("Mariam Shapira", "100%", "Managing Member"),
        ],
        col_widths=(2.5, 1.5, 3.0),
    )

    # ===================================================================
    # SECTION 4 — Operating Agreement Summary
    # ===================================================================
    add_section_heading(doc, "4", "Operating Agreement Summary")

    add_bullet(doc, "Type: Single-Member LLC")
    add_bullet(doc, "Formation date: February 1, 2018")
    add_bullet(doc, "Managing Member has full authority to:")

    add_sub_bullet(doc, "Buy, sell, lease, or encumber real property")
    add_sub_bullet(doc, "Execute contracts and agreements")
    add_sub_bullet(doc, "Manage day-to-day operations")

    add_bullet(doc, "No additional member consent required for sale of assets")
    add_bullet(doc,
               "No transfer restrictions that would impede the proposed sale")
    add_bullet(doc,
               "Dissolution provisions: At the election of the Managing Member")

    # ===================================================================
    # SECTION 5 — Good Standing Verification
    # ===================================================================
    add_section_heading(doc, "5", "Good Standing Verification")

    add_bullet(doc, "Florida Division of Corporations: Active/Current")
    add_bullet(doc, "Annual Report: Filed and current through 2026")
    add_bullet(doc, "No administrative actions pending")
    add_bullet(doc, "No dissolution proceedings")

    # ===================================================================
    # SECTION 6 — Tax Status
    # ===================================================================
    add_section_heading(doc, "6", "Tax Status")

    add_bullet(doc, "Federal EIN: 82-XXXXXXX")
    add_bullet(doc, "State Registration: Current")
    add_bullet(doc, "Property Tax Account: Current through 2025")
    add_bullet(doc, "No outstanding tax liens (per title search)")

    # ===================================================================
    # SECTION 7 — Authority to Sell
    # ===================================================================
    add_section_heading(doc, "7", "Authority to Sell")

    add_body_text(
        doc,
        "Based on review of the Operating Agreement and Florida LLC statutes "
        "(Chapter 605):",
        space_after=6,
    )

    add_bullet(
        doc,
        "Mariam Shapira, as sole Managing Member with 100% ownership interest, "
        "has full and unrestricted authority to execute a sale of the property "
        f"at {PROPERTY['address']}",
    )
    add_bullet(
        doc,
        "No additional member votes, consents, or approvals are required",
    )
    add_bullet(
        doc,
        "No right of first refusal or transfer restrictions apply",
    )

    # ===================================================================
    # CERTIFICATION
    # ===================================================================
    # Divider
    add_paragraph(doc, "", space_after=6, space_before=12)

    cert_heading_p = doc.add_paragraph()
    cert_run = cert_heading_p.add_run("Certification")
    set_run_font(cert_run, bold=True, size=12, color=NAVY_RGB)
    cert_heading_p.paragraph_format.space_before = Pt(16)
    cert_heading_p.paragraph_format.space_after = Pt(8)

    # Certification text block
    cert_text = (
        f"I, Mariam Shapira, as the sole Managing Member of "
        f"{PROPERTY['owner_entity']}, certify that the information contained "
        f"in this summary is true and correct to the best of my knowledge "
        f"as of [DATE]."
    )
    cert_p = doc.add_paragraph()
    cert_r = cert_p.add_run(f"\u201c{cert_text}\u201d")
    set_run_font(cert_r, size=11, italic=True)
    cert_p.paragraph_format.space_after = Pt(18)
    cert_p.paragraph_format.left_indent = Inches(0.5)

    # Signature line
    add_paragraph(doc, "", space_after=6, space_before=6)

    sig_line = doc.add_paragraph()
    sig_run = sig_line.add_run("_" * 45)
    set_run_font(sig_run, size=11)
    sig_line.paragraph_format.space_after = Pt(2)

    name_p = doc.add_paragraph()
    name_r = name_p.add_run("Mariam Shapira, Managing Member")
    set_run_font(name_r, size=11, bold=True)
    name_p.paragraph_format.space_after = Pt(2)

    entity_p = doc.add_paragraph()
    entity_r = entity_p.add_run(PROPERTY["owner_entity"])
    set_run_font(entity_r, size=11)
    entity_p.paragraph_format.space_after = Pt(6)

    date_p = doc.add_paragraph()
    date_r = date_p.add_run("Date: _______________")
    set_run_font(date_r, size=11)
    date_p.paragraph_format.space_after = Pt(6)

    # ===================================================================
    # FOOTER NOTE
    # ===================================================================
    add_paragraph(doc, "", space_after=12, space_before=18)

    footer_text = (
        "This Entity Documentation Summary has been prepared for the purpose of "
        "facilitating the sale of the property located at "
        f"{PROPERTY['address']}. It is intended for use by the parties to the "
        "transaction and their legal counsel. This document does not constitute "
        "legal advice."
    )
    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(footer_text)
    set_run_font(footer_r, size=9, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    footer_p.paragraph_format.space_after = Pt(6)

    return doc


# ===========================================================================
# Main
# ===========================================================================

def main():
    doc = build_entity_summary()
    filepath = output_path("12_entity_summary.docx")
    doc.save(filepath)
    size_kb = os.path.getsize(filepath) / 1024
    print(f"Created 12_entity_summary.docx at {filepath}")
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
