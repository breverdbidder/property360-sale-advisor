"""
Generate 04_sample_lease_unit201.docx — Florida Residential Lease Agreement
for Unit 201 (David & Ana Rodriguez) at Palm Bay Palms Apartments.
"""
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from data import *

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False,
                 color=None):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure the font name is respected by Word (set eastAsia hint)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text, style=None, bold=False, size=11,
                  alignment=None, space_after=6, space_before=0,
                  font_name="Times New Roman"):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    if style:
        p.style = style
    run = p.add_run(text)
    set_run_font(run, name=font_name, size=size, bold=bold)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_section_heading(doc, number, title):
    """Add a numbered section heading (e.g., '1. PARTIES')."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number}. {title}")
    set_run_font(run, bold=True, size=11)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    return p


def add_subsection(doc, number, text, indent=0.5):
    """Add a numbered subsection paragraph with indent."""
    p = doc.add_paragraph()
    run = p.add_run(f"{number} {text}")
    set_run_font(run, size=11)
    pf = p.paragraph_format
    pf.left_indent = Inches(indent)
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)
    return p


def add_body_text(doc, text, indent=0.0, space_after=4):
    """Add a regular body text paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    pf = p.paragraph_format
    if indent > 0:
        pf.left_indent = Inches(indent)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(2)
    return p


def add_signature_block(doc, label, name, extra_line=None):
    """Add a signature line with label."""
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    run = p.add_run("_" * 50 + "    Date: " + "_" * 12)
    set_run_font(run, size=11)
    pf = p.paragraph_format
    pf.space_after = Pt(2)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(label)
    set_run_font(run2, size=11, bold=True)
    pf2 = p2.paragraph_format
    pf2.space_after = Pt(0)

    if name:
        p3 = doc.add_paragraph()
        run3 = p3.add_run(name)
        set_run_font(run3, size=11)
        pf3 = p3.paragraph_format
        pf3.space_after = Pt(0)

    if extra_line:
        p4 = doc.add_paragraph()
        run4 = p4.add_run(extra_line)
        set_run_font(run4, size=11, italic=True)
        pf4 = p4.paragraph_format
        pf4.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Lease data from data.py (Unit 201)
# ---------------------------------------------------------------------------

# Find unit 201 data
UNIT_201 = None
for u in UNITS:
    if u[U_NUM] == "201":
        UNIT_201 = u
        break

UNIT_NUM = UNIT_201[U_NUM]
UNIT_TYPE = UNIT_201[U_TYPE]
UNIT_SF = UNIT_201[U_SF]
TENANT_NAME = UNIT_201[U_TENANT]
LEASE_START = UNIT_201[U_LEASE_START]
LEASE_END = UNIT_201[U_LEASE_END]
MONTHLY_RENT = UNIT_201[U_RENT]
SECURITY_DEP = UNIT_201[U_DEPOSIT]

LANDLORD_ENTITY = PROPERTY["owner_entity"]
PROP_ADDRESS = PROPERTY["address"]
YEAR_BUILT = PROPERTY["year_built"]


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_lease():
    doc = Document()

    # -- Default font for the whole document --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # =====================================================================
    # TITLE
    # =====================================================================
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("FLORIDA RESIDENTIAL LEASE AGREEMENT")
    set_run_font(title_run, size=16, bold=True)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run("(Governed by Florida Statutes Chapter 83, Part II — Residential Tenancies)")
    set_run_font(sub_run, size=10, italic=True)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)

    # =====================================================================
    # 1. PARTIES
    # =====================================================================
    add_section_heading(doc, "1", "PARTIES")

    add_body_text(doc,
        f'This Lease Agreement ("Lease") is entered into as of October 1, 2025, '
        f'by and between:')

    add_body_text(doc,
        f'LANDLORD: {LANDLORD_ENTITY}, a Florida limited liability company '
        f'("Landlord")', indent=0.5)

    add_body_text(doc,
        f'TENANT: David Rodriguez and Ana Rodriguez, jointly and severally '
        f'("Tenant")', indent=0.5)

    # =====================================================================
    # 2. PROPERTY
    # =====================================================================
    add_section_heading(doc, "2", "PROPERTY")

    add_body_text(doc,
        f'Landlord hereby leases to Tenant the following described premises '
        f'("Premises"):')

    add_body_text(doc,
        f'Address: 2750 Malabar Road SE, Unit {UNIT_NUM}, Palm Bay, FL 32907',
        indent=0.5)
    add_body_text(doc,
        f'Unit: {UNIT_NUM} — {UNIT_TYPE}, {UNIT_SF:,} square feet',
        indent=0.5)
    add_body_text(doc,
        'The Premises shall be used exclusively as a private residential '
        'dwelling and for no other purpose.',
        indent=0.5)

    # =====================================================================
    # 3. TERM
    # =====================================================================
    add_section_heading(doc, "3", "LEASE TERM")

    add_body_text(doc,
        f'The term of this Lease shall be twelve (12) months, commencing on '
        f'October 1, 2025, and ending on September 30, 2026 ("Lease Term").')

    add_body_text(doc,
        'Upon expiration of the Lease Term, this Lease shall automatically '
        'convert to a month-to-month tenancy under the same terms and conditions, '
        'unless either party provides written notice of termination at least '
        'fifteen (15) days prior to the end of any monthly period, pursuant to '
        'Florida Statute 83.57.')

    # =====================================================================
    # 4. RENT
    # =====================================================================
    add_section_heading(doc, "4", "RENT")

    add_subsection(doc, "4.1",
        f'Monthly Rent. Tenant agrees to pay Landlord the sum of '
        f'${MONTHLY_RENT:,.2f} per month ("Monthly Rent").')

    add_subsection(doc, "4.2",
        'Due Date. Rent shall be due and payable on the first (1st) day of '
        'each calendar month during the Lease Term.')

    add_subsection(doc, "4.3",
        'Late Payment. Rent not received by the fifth (5th) day of the month '
        'shall be considered late. A late fee of $50.00 shall be assessed for '
        'any rent payment received after the fifth (5th) day.')

    add_subsection(doc, "4.4",
        'Payment Method. Rent shall be paid by check, money order, cashier\'s '
        'check, or electronic funds transfer to the account designated by '
        'Landlord. Cash payments shall not be accepted unless mutually agreed '
        'upon in writing.')

    add_subsection(doc, "4.5",
        'Returned Checks. A fee of $35.00 shall be charged for any returned '
        'check. After two (2) returned checks, Landlord may require all future '
        'payments in certified funds.')

    # =====================================================================
    # 5. SECURITY DEPOSIT
    # =====================================================================
    add_section_heading(doc, "5", "SECURITY DEPOSIT (FL STATUTE 83.49)")

    add_subsection(doc, "5.1",
        f'Amount. Tenant has deposited the sum of ${SECURITY_DEP:,.2f} as a '
        f'security deposit ("Security Deposit").')

    add_subsection(doc, "5.2",
        'Holding Method. The Security Deposit shall be held in a separate '
        'non-interest-bearing account at First Southern Bank, Palm Bay, FL, '
        'pursuant to Florida Statute 83.49.')

    add_subsection(doc, "5.3",
        'Notice of Deposit. Within thirty (30) days of receipt of the Security '
        'Deposit, Landlord shall notify Tenant in writing of the name and '
        'address of the depository and whether the deposit is held in a '
        'separate account, commingled account, or posted as surety bond, '
        'as required by FL Statute 83.49(1).')

    add_subsection(doc, "5.4",
        'Return of Deposit. Upon termination of this Lease and surrender of '
        'the Premises:\n'
        '    (a) If Landlord does not intend to impose a claim on the Security '
        'Deposit, the deposit shall be returned to Tenant within fifteen (15) '
        'days.\n'
        '    (b) If Landlord intends to impose a claim against the Security '
        'Deposit, Landlord shall provide written notice by certified mail to '
        'Tenant\'s last known address within thirty (30) days after surrender '
        'of the Premises, specifying the reason for the claim, pursuant to FL '
        'Statute 83.49(3).\n'
        '    (c) Tenant shall have fifteen (15) days from receipt of the claim '
        'notice to object. If Tenant fails to object within such period, '
        'Landlord may deduct the claim from the deposit.')

    add_subsection(doc, "5.5",
        'Permitted Deductions. The Security Deposit may be applied to unpaid '
        'rent, damages beyond normal wear and tear, early termination charges, '
        'and other charges authorized by this Lease or Florida law.')

    # =====================================================================
    # 6. UTILITIES
    # =====================================================================
    add_section_heading(doc, "6", "UTILITIES AND SERVICES")

    add_subsection(doc, "6.1",
        'Tenant Responsibilities. Tenant shall be responsible for establishing '
        'accounts and paying for the following utilities:\n'
        '    - Electric service\n'
        '    - Water service')

    add_subsection(doc, "6.2",
        'Landlord Responsibilities. Landlord shall provide and pay for the '
        'following services:\n'
        '    - Trash collection\n'
        '    - Sewer service\n'
        '    - Common area maintenance (including exterior lighting, '
        'landscaping, and shared hallways)')

    add_subsection(doc, "6.3",
        'Utility Transfer. Tenant shall arrange for utility service in '
        'Tenant\'s name to commence on the first day of the Lease Term. '
        'Failure to maintain utility services may constitute a material '
        'breach of this Lease.')

    # =====================================================================
    # 7. PET POLICY
    # =====================================================================
    add_section_heading(doc, "7", "PET POLICY")

    add_subsection(doc, "7.1",
        'No Pets. No pets shall be kept on or about the Premises without '
        'the prior written consent of Landlord via a separate Pet Addendum.')

    add_subsection(doc, "7.2",
        'Pet Addendum Terms. If a pet is approved in writing by Landlord, '
        'Tenant shall pay:\n'
        '    - A non-refundable pet deposit of $250.00\n'
        '    - Monthly pet rent of $25.00 per month\n'
        'Specific pet terms, including breed and weight restrictions, shall '
        'be set forth in the Pet Addendum.')

    add_subsection(doc, "7.3",
        'Service and Support Animals. Service animals and emotional support '
        'animals are exempt from the pet policy, pet deposit, and pet rent '
        'in accordance with the Fair Housing Act (42 U.S.C. 3604) and '
        'applicable Florida law. Tenant must provide appropriate documentation '
        'upon request.')

    # =====================================================================
    # 8. NOTICE REQUIREMENTS (FL 83.56)
    # =====================================================================
    add_section_heading(doc, "8", "NOTICE REQUIREMENTS (FL STATUTE 83.56)")

    add_subsection(doc, "8.1",
        'Material Noncompliance — Non-Monetary. If Tenant materially fails '
        'to comply with the terms of this Lease (other than failure to pay '
        'rent), Landlord shall deliver a written notice specifying the '
        'noncompliance and allowing seven (7) days for Tenant to cure the '
        'noncompliance, pursuant to FL Statute 83.56(2).')

    add_subsection(doc, "8.2",
        'Material Noncompliance — Monetary. If Tenant fails to pay rent when '
        'due, Landlord shall deliver a three (3) day written notice demanding '
        'payment of rent or possession of the Premises, pursuant to FL Statute '
        '83.56(3). The three-day period excludes weekends and legal holidays.')

    add_subsection(doc, "8.3",
        'Lease Termination. For month-to-month tenancies, either party may '
        'terminate the tenancy by providing at least fifteen (15) days\' '
        'written notice prior to the end of any monthly period, pursuant '
        'to FL Statute 83.57(3).')

    # =====================================================================
    # 9. MOLD DISCLOSURE (FL 404.056)
    # =====================================================================
    add_section_heading(doc, "9", "MOLD DISCLOSURE (FL STATUTE 404.056)")

    add_body_text(doc,
        'Landlord has no knowledge of the presence of mold or mold-producing '
        'conditions in the Premises as of the date of this Lease. Tenant is '
        'hereby advised that mold may exist in any indoor environment and can '
        'affect health.')

    add_body_text(doc,
        'Tenant agrees to:\n'
        '    (a) Promptly notify Landlord in writing of any water intrusion, '
        'leaks, moisture accumulation, or visible mold growth;\n'
        '    (b) Maintain the Premises in a manner that prevents excess '
        'moisture accumulation;\n'
        '    (c) Use exhaust fans and ensure adequate ventilation;\n'
        '    (d) Promptly report any malfunctioning air conditioning or '
        'plumbing to Landlord.',
        indent=0.5)

    # =====================================================================
    # 10. RADON GAS DISCLOSURE
    # =====================================================================
    add_section_heading(doc, "10", "RADON GAS DISCLOSURE")

    add_body_text(doc,
        'RADON GAS: Radon is a naturally occurring radioactive gas that, '
        'when it has accumulated in a building in sufficient quantities, may '
        'present health risks to persons who are exposed to it over time. '
        'Levels of radon that exceed federal and state guidelines have been '
        'found in buildings in Florida. Additional information regarding radon '
        'and radon testing may be obtained from your county health department.')

    add_body_text(doc,
        'This disclosure is provided in accordance with Section 404.056(5), '
        'Florida Statutes.',
        indent=0.5)

    # =====================================================================
    # 11. LEAD-BASED PAINT DISCLOSURE
    # =====================================================================
    add_section_heading(doc, "11", "LEAD-BASED PAINT DISCLOSURE")

    add_body_text(doc,
        f'The Property was built in {YEAR_BUILT}. Because the Property was '
        f'constructed after 1978, lead-based paint disclosure is not required '
        f'under the Residential Lead-Based Paint Hazard Reduction Act of 1992 '
        f'(42 U.S.C. 4852d). However, Landlord provides the following '
        f'disclosure for completeness:')

    add_body_text(doc,
        '    (a) Landlord has no knowledge of lead-based paint and/or '
        'lead-based paint hazards in the Premises.\n'
        '    (b) Landlord has no reports or records pertaining to lead-based '
        'paint and/or lead-based paint hazards in the Premises.',
        indent=0.5)

    # =====================================================================
    # 12. MAINTENANCE AND REPAIRS
    # =====================================================================
    add_section_heading(doc, "12", "MAINTENANCE AND REPAIRS")

    add_subsection(doc, "12.1",
        'Landlord Obligations. Landlord shall maintain the Premises in '
        'compliance with applicable building, housing, and health codes; '
        'maintain the roof, structural components, plumbing, electrical, '
        'and HVAC systems; and comply with all obligations imposed by '
        'FL Statute 83.51.')

    add_subsection(doc, "12.2",
        'Tenant Obligations. Tenant shall keep the Premises clean and '
        'sanitary; use all fixtures, appliances, and equipment in a '
        'reasonable manner; not destroy or damage the Premises; and '
        'comply with all obligations imposed by FL Statute 83.52.')

    # =====================================================================
    # 13. NO SMOKING
    # =====================================================================
    add_section_heading(doc, "13", "SMOKING POLICY")

    add_body_text(doc,
        'Smoking (including e-cigarettes and vaping) is strictly prohibited '
        'in all units, common areas, hallways, stairwells, and within '
        'twenty-five (25) feet of any building entrance. Violation of this '
        'policy shall constitute a material breach of this Lease.')

    # =====================================================================
    # 14. QUIET ENJOYMENT
    # =====================================================================
    add_section_heading(doc, "14", "QUIET ENJOYMENT AND NOISE POLICY")

    add_body_text(doc,
        'Quiet hours shall be observed from 10:00 PM to 8:00 AM daily. '
        'Tenant shall not engage in or permit any activity that unreasonably '
        'disturbs the quiet enjoyment of other tenants. Repeated noise '
        'violations shall constitute grounds for lease termination pursuant '
        'to FL Statute 83.56.')

    # =====================================================================
    # 15. RENTER'S INSURANCE
    # =====================================================================
    add_section_heading(doc, "15", "RENTER'S INSURANCE")

    add_body_text(doc,
        'Tenant is responsible for obtaining and maintaining renter\'s '
        'insurance throughout the Lease Term with a recommended minimum '
        'liability coverage of $100,000. Landlord\'s insurance does not '
        'cover Tenant\'s personal property or liability. Tenant is strongly '
        'encouraged to obtain a policy prior to move-in.')

    # =====================================================================
    # 16. ALTERATIONS
    # =====================================================================
    add_section_heading(doc, "16", "ALTERATIONS")

    add_body_text(doc,
        'Tenant shall not make any alterations, additions, or improvements '
        'to the Premises without the prior written consent of Landlord. '
        'Any approved alterations shall become the property of Landlord '
        'upon termination of this Lease, unless otherwise agreed in writing.')

    # =====================================================================
    # 17. RIGHT OF ENTRY
    # =====================================================================
    add_section_heading(doc, "17", "RIGHT OF ENTRY (FL STATUTE 83.53)")

    add_body_text(doc,
        'Landlord may enter the Premises with at least twelve (12) hours\' '
        'notice to Tenant for the following purposes, pursuant to FL Statute '
        '83.53:')

    add_body_text(doc,
        '    (a) To inspect the Premises;\n'
        '    (b) To make necessary or agreed-upon repairs, decorations, '
        'alterations, or improvements;\n'
        '    (c) To supply agreed-upon services;\n'
        '    (d) To exhibit the Premises to prospective or actual purchasers, '
        'mortgagees, tenants, workers, or contractors.',
        indent=0.5)

    add_body_text(doc,
        'Landlord may enter the Premises at any time without notice in the '
        'event of an emergency. Entry shall be made at reasonable times, '
        'except in the case of emergency.')

    # =====================================================================
    # 18. DEFAULT AND REMEDIES
    # =====================================================================
    add_section_heading(doc, "18", "DEFAULT AND REMEDIES")

    add_subsection(doc, "18.1",
        'Tenant Default. If Tenant fails to pay rent or otherwise materially '
        'breaches any term of this Lease, Landlord may pursue remedies '
        'available under Florida Statutes Chapter 83, Part II, including '
        'but not limited to termination of this Lease and recovery of '
        'possession of the Premises.')

    add_subsection(doc, "18.2",
        'Landlord Default. If Landlord materially fails to comply with the '
        'requirements of FL Statute 83.51 or the material provisions of this '
        'Lease, Tenant may pursue remedies as provided by FL Statute 83.60.')

    add_subsection(doc, "18.3",
        'Attorney\'s Fees. In any action arising under this Lease, the '
        'prevailing party shall be entitled to recover reasonable attorney\'s '
        'fees and costs, pursuant to FL Statute 83.48.')

    # =====================================================================
    # 19. GENERAL PROVISIONS
    # =====================================================================
    add_section_heading(doc, "19", "GENERAL PROVISIONS")

    add_subsection(doc, "19.1",
        'Entire Agreement. This Lease constitutes the entire agreement '
        'between the parties and supersedes all prior negotiations, '
        'representations, or agreements.')

    add_subsection(doc, "19.2",
        'Amendments. This Lease may not be modified or amended except by '
        'a written instrument signed by both Landlord and Tenant.')

    add_subsection(doc, "19.3",
        'Severability. If any provision of this Lease is found to be invalid '
        'or unenforceable, the remaining provisions shall continue in full '
        'force and effect.')

    add_subsection(doc, "19.4",
        'Governing Law. This Lease shall be governed by and construed in '
        'accordance with the laws of the State of Florida.')

    add_subsection(doc, "19.5",
        'Joint and Several Liability. If Tenant consists of more than one '
        'person, each person shall be jointly and severally liable for all '
        'obligations of Tenant under this Lease.')

    add_subsection(doc, "19.6",
        'Notices. All notices required or permitted under this Lease shall '
        'be in writing and shall be delivered personally, sent by certified '
        'mail (return receipt requested), or sent by overnight courier to '
        'the addresses set forth herein or to such other addresses as the '
        'parties may designate in writing.')

    # =====================================================================
    # SIGNATURES
    # =====================================================================
    add_paragraph(doc, "", space_after=12)

    sig_header = doc.add_paragraph()
    sig_header_run = sig_header.add_run("SIGNATURES")
    set_run_font(sig_header_run, size=12, bold=True)
    sig_header.paragraph_format.space_before = Pt(18)
    sig_header.paragraph_format.space_after = Pt(6)

    add_body_text(doc,
        'IN WITNESS WHEREOF, the parties have executed this Lease Agreement '
        'as of the date first written above.')

    # Landlord signature
    add_signature_block(
        doc,
        "LANDLORD:",
        LANDLORD_ENTITY,
        "By: Mariam Shapira, Managing Member"
    )

    # Tenant 1 signature
    add_signature_block(
        doc,
        "TENANT:",
        "David Rodriguez",
        None
    )

    # Tenant 2 signature
    add_signature_block(
        doc,
        "TENANT:",
        "Ana Rodriguez",
        None
    )

    # =====================================================================
    # NOTARY BLOCK
    # =====================================================================
    doc.add_page_break()

    notary_header = doc.add_paragraph()
    notary_run = notary_header.add_run("NOTARY ACKNOWLEDGMENT")
    set_run_font(notary_run, size=12, bold=True)
    notary_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notary_header.paragraph_format.space_after = Pt(12)

    add_body_text(doc,
        'STATE OF FLORIDA')
    add_body_text(doc,
        'COUNTY OF BREVARD')

    add_body_text(doc, "", space_after=6)

    add_body_text(doc,
        'Before me, the undersigned notary public, on this ______ day of '
        '______________, 20____, personally appeared:')

    add_body_text(doc, "", space_after=6)

    add_body_text(doc,
        'Name(s): _____________________________________________________________',
        indent=0.5)

    add_body_text(doc, "", space_after=6)

    add_body_text(doc,
        'who is/are personally known to me or has/have produced '
        '_________________________________ as identification, and who did '
        'acknowledge before me that he/she/they executed the foregoing '
        'instrument for the purposes therein expressed.')

    add_body_text(doc, "", space_after=18)

    # Notary signature line
    p_notary_sig = doc.add_paragraph()
    run_notary = p_notary_sig.add_run(
        "_" * 45 + "\n"
        "Notary Public, State of Florida\n"
        "My Commission Expires: _______________\n"
        "Commission No.: _____________________"
    )
    set_run_font(run_notary, size=11)
    p_notary_sig.paragraph_format.space_before = Pt(12)

    # Notary seal placeholder
    add_body_text(doc, "", space_after=12)

    seal_p = doc.add_paragraph()
    seal_run = seal_p.add_run("[NOTARY SEAL]")
    set_run_font(seal_run, size=10, italic=True)
    seal_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


# ===========================================================================
# Main
# ===========================================================================

def main():
    doc = build_lease()
    filepath = output_path("04_sample_lease_unit201.docx")
    doc.save(filepath)
    size_kb = os.path.getsize(filepath) / 1024
    print(f"Created 04_sample_lease_unit201.docx at {filepath}")
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
