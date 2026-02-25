"""
Generate 08_loi_template.docx — Letter of Intent (LOI) for the purchase of
Palm Bay Palms Apartments, 2750 Malabar Road SE, Palm Bay, FL 32907.
Standard commercial real-estate LOI with 15 sections, signature blocks,
and fill-in blanks for buyer-specific details.
"""
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from data import *

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False,
                 color=None, underline=False):
    """Apply font settings to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if underline:
        run.underline = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure the font name is respected by Word (set eastAsia hint)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc, text, bold=False, size=11, alignment=None,
                  space_after=6, space_before=0, font_name="Times New Roman"):
    """Add a paragraph with consistent formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, name=font_name, size=size, bold=bold)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_mixed_paragraph(doc, runs_spec, alignment=None, space_after=6,
                        space_before=0):
    """Add a paragraph with mixed formatting.
    runs_spec: list of dicts with keys 'text' and optional 'bold', 'italic',
               'underline', 'size'.
    """
    p = doc.add_paragraph()
    for spec in runs_spec:
        run = p.add_run(spec["text"])
        set_run_font(
            run,
            size=spec.get("size", 11),
            bold=spec.get("bold", False),
            italic=spec.get("italic", False),
            underline=spec.get("underline", False),
        )
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p


def add_section(doc, number, title, body_lines, space_before_heading=14):
    """Add a numbered section with bold header and body paragraphs.
    body_lines: list of strings (each becomes its own paragraph).
    """
    # Section heading
    p_head = doc.add_paragraph()
    run_head = p_head.add_run(f"Section {number} \u2014 {title}")
    set_run_font(run_head, bold=True, size=11)
    pf_head = p_head.paragraph_format
    pf_head.space_before = Pt(space_before_heading)
    pf_head.space_after = Pt(4)

    # Body text
    for line in body_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, size=11)
        pf = p.paragraph_format
        pf.space_after = Pt(4)
        pf.space_before = Pt(2)


# ---------------------------------------------------------------------------
# Build the LOI document
# ---------------------------------------------------------------------------

def build_loi():
    doc = Document()

    # -- Default font for the whole document --------------------------------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)

    # Set 1" margins all around
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===================================================================
    # TITLE
    # ===================================================================
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("LETTER OF INTENT")
    set_run_font(title_run, size=14, bold=True)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(2)

    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run("Commercial Real Estate Purchase")
    set_run_font(subtitle_run, size=14, bold=True)
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(18)

    # ===================================================================
    # DATE
    # ===================================================================
    add_paragraph(doc, "Date: [DATE]", bold=False, space_after=12)

    # ===================================================================
    # TO / FROM / RE
    # ===================================================================
    # To block
    add_mixed_paragraph(doc, [
        {"text": "To:       ", "bold": True},
        {"text": f'{PROPERTY["owner_entity"]} ("Seller")'},
    ], space_after=2)
    add_paragraph(doc, "          Attn: Mariam Shapira, Managing Member",
                  space_after=12, space_before=0)

    # From block
    add_mixed_paragraph(doc, [
        {"text": "From:   ", "bold": True},
        {"text": '[BUYER NAME] ("Buyer")'},
    ], space_after=12)

    # Re block
    add_mixed_paragraph(doc, [
        {"text": "Re:       ", "bold": True},
        {"text": f"Letter of Intent to Purchase \u2014 {PROPERTY['address']}"},
    ], space_after=14)

    # ===================================================================
    # OPENING PARAGRAPH
    # ===================================================================
    add_paragraph(
        doc,
        "This Letter of Intent sets forth the basic terms upon which Buyer "
        "proposes to purchase the property described herein. Except for the "
        "provisions regarding Confidentiality and Exclusivity, this LOI is "
        "non-binding.",
        space_after=10,
    )

    # ===================================================================
    # SECTION 1 — PROPERTY
    # ===================================================================
    add_section(doc, "1", "Property", [
        f"Address: {PROPERTY['address']}",
        f"Type: {PROPERTY['property_type']}",
        "Legal: Lot 142, Block 5, Palm Bay Unit 37, Brevard County, FL",
    ])

    # ===================================================================
    # SECTION 2 — PURCHASE PRICE
    # ===================================================================
    add_section(doc, "2", "Purchase Price", [
        "$__________ (fill-in blank)",
    ])

    # ===================================================================
    # SECTION 3 — EARNEST MONEY
    # ===================================================================
    add_section(doc, "3", "Earnest Money", [
        "$__________ (to be deposited within 3 business days of execution). "
        "Earnest money becomes non-refundable (\u201chard\u201d) after the "
        "expiration of the 15-day inspection period.",
    ])

    # ===================================================================
    # SECTION 4 — INSPECTION PERIOD
    # ===================================================================
    add_section(doc, "4", "Inspection Period", [
        "15 business days from the Effective Date. Buyer may terminate for "
        "any reason during this period and receive full refund of earnest money.",
    ])

    # ===================================================================
    # SECTION 5 — FINANCING CONTINGENCY
    # ===================================================================
    add_section(doc, "5", "Financing Contingency", [
        "30 days from the Effective Date. Buyer shall apply for financing "
        "within 5 business days and diligently pursue loan approval.",
    ])

    # ===================================================================
    # SECTION 6 — CLOSING DATE
    # ===================================================================
    add_section(doc, "6", "Closing Date", [
        "45 days from execution of this LOI, or 10 business days after "
        "satisfaction of all contingencies, whichever is later.",
    ])

    # ===================================================================
    # SECTION 7 — INCLUDED IN SALE
    # ===================================================================
    add_section(doc, "7", "Included in Sale", [
        "All fixtures, appliances, laundry equipment, HVAC systems, and "
        "other property attached to or used in connection with the Property, "
        "including all tenant security deposits and existing service contracts "
        "(unless terminated per Section 4).",
    ])

    # ===================================================================
    # SECTION 8 — EXCLUDED
    # ===================================================================
    add_section(doc, "8", "Excluded", [
        "Personal property of tenants, Seller\u2019s proprietary business "
        "records not directly related to property operations.",
    ])

    # ===================================================================
    # SECTION 9 — 1031 EXCHANGE
    # ===================================================================
    add_section(doc, "9", "1031 Exchange", [
        "Buyer and/or Seller may structure this transaction as a tax-deferred "
        "exchange under IRC Section 1031. Each party agrees to reasonably "
        "cooperate with the other\u2019s exchange, provided such cooperation "
        "does not delay closing or impose additional cost on the non-exchanging "
        "party.",
    ])

    # ===================================================================
    # SECTION 10 — DUE DILIGENCE MATERIALS
    # ===================================================================
    add_section(doc, "10", "Due Diligence Materials", [
        "Within 5 business days of the Effective Date, Seller shall provide: "
        "current rent roll, trailing 12-month P&L, copies of all leases, "
        "insurance policy, tax bills, vendor contracts, building permits, and "
        "such other documents as Buyer may reasonably request.",
    ])

    # ===================================================================
    # SECTION 11 — REPRESENTATIONS
    # ===================================================================
    add_section(doc, "11", "Representations", [
        "Seller represents that to Seller\u2019s knowledge:",
        "    (a) there are no pending or threatened legal actions;",
        "    (b) all material defects have been disclosed;",
        "    (c) all leases provided are true and complete copies.",
    ])

    # ===================================================================
    # SECTION 12 — CONFIDENTIALITY
    # ===================================================================
    p12_head = doc.add_paragraph()
    run12_num = p12_head.add_run("Section 12 \u2014 Confidentiality")
    set_run_font(run12_num, bold=True, size=11)
    pf12 = p12_head.paragraph_format
    pf12.space_before = Pt(14)
    pf12.space_after = Pt(4)

    p12_binding = doc.add_paragraph()
    r12_bind = p12_binding.add_run("BINDING.")
    set_run_font(r12_bind, bold=True, size=11)
    r12_text = p12_binding.add_run(
        " The parties agree to keep the terms of this LOI and all information "
        "exchanged confidential, except as required by law or as necessary to "
        "consummate the transaction (attorneys, lenders, accountants)."
    )
    set_run_font(r12_text, size=11)
    p12_binding.paragraph_format.space_after = Pt(4)

    # ===================================================================
    # SECTION 13 — EXCLUSIVITY
    # ===================================================================
    p13_head = doc.add_paragraph()
    run13_num = p13_head.add_run("Section 13 \u2014 Exclusivity")
    set_run_font(run13_num, bold=True, size=11)
    pf13 = p13_head.paragraph_format
    pf13.space_before = Pt(14)
    pf13.space_after = Pt(4)

    p13_binding = doc.add_paragraph()
    r13_bind = p13_binding.add_run("BINDING.")
    set_run_font(r13_bind, bold=True, size=11)
    r13_text = p13_binding.add_run(
        " For a period of 30 days from execution, Seller agrees not to market "
        "the Property, solicit offers, or negotiate with other potential "
        "purchasers."
    )
    set_run_font(r13_text, size=11)
    p13_binding.paragraph_format.space_after = Pt(4)

    # ===================================================================
    # SECTION 14 — BINDING EFFECT
    # ===================================================================
    add_section(doc, "14", "Binding Effect", [
        "This LOI is non-binding except for Sections 12 (Confidentiality) and "
        "13 (Exclusivity), which shall survive regardless of whether a "
        "definitive Purchase Agreement is executed.",
    ])

    # ===================================================================
    # SECTION 15 — EXPIRATION
    # ===================================================================
    add_section(doc, "15", "Expiration", [
        "This LOI shall expire if not executed by both parties within "
        "10 business days of the date first written above.",
    ])

    # ===================================================================
    # SIGNATURE BLOCKS
    # ===================================================================
    add_paragraph(doc, "", space_after=6, space_before=24)

    # ----- BUYER SIGNATURE -----
    buyer_label = doc.add_paragraph()
    r_buyer = buyer_label.add_run("BUYER:")
    set_run_font(r_buyer, bold=True, size=11)
    buyer_label.paragraph_format.space_before = Pt(18)
    buyer_label.paragraph_format.space_after = Pt(6)

    add_paragraph(doc, "___________________________________", space_after=2)
    add_paragraph(doc, "[BUYER NAME]", space_after=8)

    add_paragraph(doc, "By: ________________________________", space_after=2)
    add_paragraph(doc, "Title: _____________________________", space_after=2)
    add_paragraph(doc, "Date: ______________________________", space_after=18)

    # ----- SELLER SIGNATURE -----
    seller_label = doc.add_paragraph()
    r_seller = seller_label.add_run("SELLER:")
    set_run_font(r_seller, bold=True, size=11)
    seller_label.paragraph_format.space_before = Pt(18)
    seller_label.paragraph_format.space_after = Pt(6)

    add_paragraph(doc, "___________________________________", space_after=2)
    add_paragraph(doc, f"{PROPERTY['owner_entity']}", space_after=8)

    add_paragraph(doc, "By: Mariam Shapira, Managing Member", space_after=2)
    add_paragraph(doc, "Date: ______________________________", space_after=6)

    return doc


# ===========================================================================
# Main
# ===========================================================================

def main():
    doc = build_loi()
    filepath = output_path("08_loi_template.docx")
    doc.save(filepath)
    size_kb = os.path.getsize(filepath) / 1024
    print(f"Created 08_loi_template.docx at {filepath}")
    print(f"File size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
